"""知识库服务（LLM-WIKI 式：Markdown 文件夹即知识库，多级目录组织）。

设计理念（对齐 Karpathy LLM-WIKI）：
- 知识库根目录默认为 backend/knowledge/，可在「设置 -> 知识库」中修改路径
  （持久化到 config/kb.json 并写入 KB_ROOT 环境变量，立即生效，无需重启）；
- 文件夹 = 真实目录（支持多级嵌套），文档以 Markdown 形式直接存放在目录中，
  LLM 或人工可直接读取/检索；上传的 PDF / Word / TXT 等原始文件解析为
  Markdown 落盘（「编译」步骤），原始文件保留在 raw/ 供下载；
- 文档索引集中在 meta.json（doc_id -> 元信息），目录/文件名操作与索引保持一致；
- 知识库无需权限：所有接口直接可用，不校验任何令牌。

线程安全：所有写操作经模块级锁串行化；FastAPI 多线程下安全。
"""
from __future__ import annotations

import json
import os
import re
import shutil
import threading
import uuid
from datetime import datetime
from typing import Any, Dict, Optional, Tuple

from .. import kb_settings

# 非法文件名字符（跨平台：Windows 保留字符 + 控制字符）
_INVALID = re.compile(r'[\\/:*?"<>|\x00-\x1f]')
_LOCK = threading.Lock()

ALLOWED_EXTS = {'.pdf', '.doc', '.docx', '.txt', '.md', '.markdown'}


def kb_root() -> str:
    """知识库根目录：设置动态配置 > KB_ROOT 环境变量 > 默认 backend/knowledge。"""
    root = kb_settings.get_kb_cfg().get('root_path', '').strip()
    if not root:
        root = os.environ.get('KB_ROOT', '').strip()
    if not root:
        root = kb_settings.DEFAULT_KB_PATH
    os.makedirs(root, exist_ok=True)
    return root


def _meta_path() -> str:
    return os.path.join(kb_root(), 'meta.json')


def _load_meta() -> Dict[str, Any]:
    try:
        with open(_meta_path(), 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {'docs': {}}


def _save_meta(meta: Dict[str, Any]) -> None:
    tmp = _meta_path() + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    os.replace(tmp, _meta_path())


def _now() -> str:
    return datetime.now().isoformat(timespec='seconds')


def _safe_name(name: str) -> str:
    """清洗非法文件名字符；空名回退为「未命名」。"""
    cleaned = _INVALID.sub('_', str(name or '')).strip().strip('.')
    return cleaned or '未命名'


def _norm_folder(folder: str) -> str:
    """规范化文件夹相对路径：'/' 或 '' 视为根；去除首尾 /；拒绝越权路径。"""
    f = (folder or '').strip().replace('\\', '/').strip('/')
    if f in ('', '/'):
        return ''
    parts = [p for p in f.split('/') if p and p not in ('.', '..')]
    return '/'.join(parts)


def _folder_abs(folder: str) -> str:
    root = kb_root()
    norm = _norm_folder(folder)
    p = root if not norm else os.path.join(root, *norm.split('/'))
    return p


def _folder_exists(folder: str) -> bool:
    return os.path.isdir(_folder_abs(folder))


def _doc_md_path(doc: Dict[str, Any]) -> str:
    folder = _norm_folder(doc.get('folder', ''))
    return os.path.join(_folder_abs(folder), doc.get('title', '未命名') + '.md')


# ------------------------- 文本抽取（上传文件 -> Markdown） -------------------------

def _read_txt(data: bytes) -> str:
    for enc in ('utf-8', 'gb18030', 'big5', 'latin-1'):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode('utf-8', errors='replace')


def extract_text(filename: str, data: bytes) -> Tuple[str, bool]:
    """按扩展名抽取文本，返回 (text, parse_ok)。无法解析时 parse_ok=False。"""
    ext = os.path.splitext(filename or '')[1].lower()
    if ext in ('.txt', '.md', '.markdown'):
        return _read_txt(data), True
    if ext == '.pdf':
        try:
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(data))
            pages = [page.extract_text() or '' for page in reader.pages]
            return '\n\n'.join(p.strip() for p in pages if p.strip()), True
        except Exception:
            return '', False
    if ext == '.docx':
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(' | '.join(cells))
            return '\n\n'.join(parts), True
        except Exception:
            return '', False
    if ext == '.doc':
        # 旧版 Word 二进制格式：无纯 Python 可靠解析器，保留原始文件、标记未解析
        return '', False
    return '', False


# ------------------------- 文件夹操作 -------------------------

def create_folder(name: str, parent: str = '') -> Dict[str, Any]:
    with _LOCK:
        parent = _norm_folder(parent)
        if parent and not _folder_exists(parent):
            return {'ok': False, 'error': '父文件夹不存在'}
        safe = _safe_name(name)
        if not safe:
            return {'ok': False, 'error': '文件夹名称不能为空'}
        path = safe if not parent else parent + '/' + safe
        if _folder_exists(path):
            return {'ok': False, 'error': '同名文件夹已存在'}
        os.makedirs(_folder_abs(path), exist_ok=True)
        return {'ok': True, 'path': path}


def rename_folder(path: str, new_name: str) -> Dict[str, Any]:
    with _LOCK:
        path = _norm_folder(path)
        if not path or not _folder_exists(path):
            return {'ok': False, 'error': '文件夹不存在'}
        safe = _safe_name(new_name)
        parent = '/' .join(path.split('/')[:-1]) if '/' in path else ''
        target = safe if not parent else parent + '/' + safe
        target = _norm_folder(target)
        if target != path and _folder_exists(target):
            return {'ok': False, 'error': '目标文件夹已存在'}
        old_abs, new_abs = _folder_abs(path), _folder_abs(target)
        # 更新索引中位于该目录下的文档 folder 前缀，并随目录移动 .md 文件
        meta = _load_meta()
        prefix = path + '/'
        for doc in meta['docs'].values():
            if doc.get('folder') == path:
                doc['folder'] = target
            elif doc.get('folder', '').startswith(prefix):
                doc['folder'] = target + '/' + doc['folder'][len(prefix):]
        os.rename(old_abs, new_abs)
        _save_meta(meta)
        return {'ok': True, 'path': target}


def delete_folder(path: str) -> Dict[str, Any]:
    with _LOCK:
        path = _norm_folder(path)
        if not path:
            return {'ok': False, 'error': '不允许删除根目录'}
        abs_dir = _folder_abs(path)
        if not os.path.isdir(abs_dir):
            return {'ok': False, 'error': '文件夹不存在'}
        meta = _load_meta()
        docs = meta['docs']
        prefix = path + '/'
        removed = [did for did, d in docs.items()
                   if d.get('folder') == path or d.get('folder', '').startswith(prefix)]
        for did in removed:
            raw = docs[did].get('raw')
            if raw and os.path.isfile(os.path.join(kb_root(), raw)):
                try:
                    os.remove(os.path.join(kb_root(), raw))
                except OSError:
                    pass
            del docs[did]
        shutil.rmtree(abs_dir, ignore_errors=True)
        _save_meta(meta)
        return {'ok': True, 'removed_docs': len(removed)}


# ------------------------- 文档操作 -------------------------

def upload(filename: str, data: bytes, folder: str = '') -> Dict[str, Any]:
    with _LOCK:
        ext = os.path.splitext(filename or '')[1].lower()
        if ext not in ALLOWED_EXTS:
            return {'ok': False, 'error': f'不支持的格式 {ext or "（无扩展名）"}，支持：pdf / doc / docx / txt / md'}
        folder = _norm_folder(folder)
        if folder and not _folder_exists(folder):
            return {'ok': False, 'error': '目标文件夹不存在'}
        title = _safe_name(os.path.splitext(os.path.basename(filename))[0])
        text, parse_ok = extract_text(filename, data)
        if not parse_ok:
            return {'ok': False, 'error': '文档解析失败：无法提取文本（PDF/Word 已损坏或格式过旧，可尝试另存为 PDF/DOCX/TXT 后重传）'}
        doc_id = uuid.uuid4().hex[:12]
        md_abs = os.path.join(_folder_abs(folder), title + '.md')
        if os.path.exists(md_abs):
            md_abs = os.path.join(_folder_abs(folder), f'{title}_{doc_id[:6]}.md')
        with open(md_abs, 'w', encoding='utf-8') as f:
            f.write(text)
        raw_rel = ''
        if ext not in ('.md', '.markdown'):
            raw_dir = os.path.join(kb_root(), 'raw')
            os.makedirs(raw_dir, exist_ok=True)
            raw_name = f'{doc_id}{ext}'
            raw_rel = os.path.join('raw', raw_name)
            with open(os.path.join(kb_root(), raw_rel), 'wb') as f:
                f.write(data)
        now = _now()
        doc = {
            'id': doc_id,
            'title': title,
            'folder': folder,
            'ext': ext,
            'size': len(data),
            'chars': len(text),
            'original': os.path.basename(filename),
            'parse': True,
            'raw': raw_rel,
            'created_at': now,
            'updated_at': now,
        }
        meta = _load_meta()
        meta['docs'][doc_id] = doc
        _save_meta(meta)
        return {'ok': True, 'doc': doc}


def rename_doc(doc_id: str, new_name: str) -> Dict[str, Any]:
    with _LOCK:
        meta = _load_meta()
        docs = meta['docs']
        doc = docs.get(doc_id)
        if not doc:
            return {'ok': False, 'error': '文档不存在'}
        title = _safe_name(new_name)
        if not title:
            return {'ok': False, 'error': '文档名称不能为空'}
        if title == doc.get('title'):
            return {'ok': True, 'doc': doc}
        old_md = _doc_md_path(doc)
        doc['title'] = title
        new_md = _doc_md_path(doc)
        if os.path.exists(new_md):
            return {'ok': False, 'error': '同名文档已存在'}
        if os.path.exists(old_md):
            os.rename(old_md, new_md)
        doc['updated_at'] = _now()
        _save_meta(meta)
        return {'ok': True, 'doc': doc}


def delete_doc(doc_id: str) -> Dict[str, Any]:
    with _LOCK:
        meta = _load_meta()
        doc = meta['docs'].pop(doc_id, None)
        if not doc:
            return {'ok': False, 'error': '文档不存在'}
        md = _doc_md_path(doc)
        if os.path.exists(md):
            try:
                os.remove(md)
            except OSError:
                pass
        raw = doc.get('raw')
        if raw and os.path.isfile(os.path.join(kb_root(), raw)):
            try:
                os.remove(os.path.join(kb_root(), raw))
            except OSError:
                pass
        _save_meta(meta)
        return {'ok': True}


def get_doc(doc_id: str) -> Optional[Dict[str, Any]]:
    return _load_meta()['docs'].get(doc_id)


def doc_content(doc_id: str) -> Dict[str, Any]:
    doc = get_doc(doc_id)
    if not doc:
        return {'ok': False, 'error': '文档不存在'}
    md = _doc_md_path(doc)
    try:
        with open(md, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        content = ''
    return {'ok': True, 'doc': doc, 'content': content}


# ------------------------- 树 / 检索 -------------------------

# 内部文件/目录（索引与原始文件存储），不进入知识库树
_INTERNAL = {'meta.json', 'raw'}
_SKIP_PREFIXES = ('.', '#', '~')


def _is_internal(name: str) -> bool:
    if name in _INTERNAL or name.endswith('.tmp'):
        return True
    return name.startswith(_SKIP_PREFIXES)


def _build_node(folder_abs: str, rel_path: str, docs: Dict[str, Any]) -> Dict[str, Any]:
    node: Dict[str, Any] = {'path': rel_path or '/', 'name': os.path.basename(folder_abs) if rel_path else '', 'docs': [], 'children': []}
    for entry in sorted(os.listdir(folder_abs), key=str.lower):
        if _is_internal(entry):
            continue
        child_abs = os.path.join(folder_abs, entry)
        child_rel = entry if not rel_path else rel_path + '/' + entry
        if os.path.isdir(child_abs):
            node['children'].append(_build_node(child_abs, child_rel, docs))
        elif entry.endswith('.md'):
            title = entry[:-3]
            # 通过索引找到对应 doc（标题 + 文件夹匹配），补充元信息
            for d in docs.values():
                if d.get('folder') == rel_path and d.get('title') == title:
                    node['docs'].append(d)
                    break
    return node


def list_tree() -> Dict[str, Any]:
    with _LOCK:
        meta = _load_meta()
        docs = meta['docs']
        root = kb_root()
        tree = _build_node(root, '', docs)
        total_chars = sum(d.get('chars', 0) for d in docs.values())
        return {
            'ok': True,
            'stats': {'folders': _count_folders(tree), 'docs': len(docs), 'chars': total_chars},
            'tree': tree,
        }


def _count_folders(node: Dict[str, Any]) -> int:
    return len(node['children']) + sum(_count_folders(c) for c in node['children'])


def search(q: str) -> Dict[str, Any]:
    """标题 / 全文检索（内容按行扫描），最多返回 50 条。"""
    with _LOCK:
        q = (q or '').strip()
        if not q:
            return {'ok': True, 'q': '', 'results': []}
        meta = _load_meta()
        hits = []
        for doc in meta['docs'].values():
            score = 0
            title = doc.get('title', '')
            if q.lower() in title.lower():
                score += 100
            body = ''
            md = _doc_md_path(doc)
            try:
                with open(md, 'r', encoding='utf-8') as f:
                    body = f.read()
            except OSError:
                body = ''
            if body:
                low = body.lower()
                idx = low.find(q.lower())
                if idx >= 0:
                    score += 20
                    start = max(0, idx - 40)
                    snippet = body[start:idx + len(q) + 60].replace('\n', ' ').strip()
                else:
                    snippet = body[:120].replace('\n', ' ').strip()
            else:
                snippet = ''
            if score:
                hits.append({'doc': doc, 'score': score, 'snippet': snippet})
        hits.sort(key=lambda h: (-h['score'], h['doc']['title']))
        return {'ok': True, 'q': q, 'results': [{'doc': h['doc'], 'snippet': h['snippet']} for h in hits[:50]]}
